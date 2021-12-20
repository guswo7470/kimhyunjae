import io
import os
import chardet
import sqlite3
import zipfile
import xlrd
from struct import pack, unpack
from collections import OrderedDict
from tika import parser
import docx

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML


class KeywordSearch:
    def __init__(self, index_folder, file_num=1):
        self.index_folder = index_folder
        if not os.path.exists(self.index_folder):
            os.makedirs(self.index_folder)
        self.index_db = os.path.join(self.index_folder, 'index.db')
        self.conn1 = sqlite3.connect(self.index_db)
        self.cur1 = self.conn1.cursor()
        self.cur1.execute("PRAGMA page_size = 65536;")
        self.cur1.execute("PRAGMA journal_mode=wal")
        self.cur1.execute("PRAGMA synchronous=OFF;")
        self.cur1.execute("CREATE TABLE IF NOT EXISTS word_set (word text, location blob)")
        self.cur1.execute("CREATE INDEX IF NOT EXISTS word_set_IDX ON word_set (word)")

        self.path_db = os.path.join(self.index_folder, 'path.db')
        self.conn2 = sqlite3.connect(self.path_db)
        self.cur2 = self.conn2.cursor()
        self.cur2.execute("PRAGMA page_size = 65536;")
        self.cur2.execute("PRAGMA journal_mode=wal")
        self.cur2.execute("PRAGMA synchronous=OFF;")
        self.cur2.execute("CREATE TABLE IF NOT EXISTS path (document_num int, path text)")
        self.cur2.execute("CREATE INDEX IF NOT EXISTS path_IDX ON path (document_num)")
        self.file_num = file_num

        self.index2_db = os.path.join(self.index_folder, 'index2.db')
        self.conn3 = sqlite3.connect(self.path_db)
        self.cur3 = self.conn3.cursor()
        self.cur3.execute("PRAGMA page_size = 65536;")
        self.cur3.execute("PRAGMA journal_mode=wal")
        self.cur3.execute("PRAGMA synchronous=OFF;")
        self.cur1.execute("CREATE TABLE IF NOT EXISTS word_set (word text, location blob)")
        self.cur1.execute("CREATE INDEX IF NOT EXISTS word_set_IDX ON word_set (word)")

        self.path2_db = os.path.join(self.index_folder, 'path2.db')
        self.conn4 = sqlite3.connect(self.path_db)
        self.cur4 = self.conn4.cursor()
        self.cur4.execute("PRAGMA page_size = 65536;")
        self.cur4.execute("PRAGMA journal_mode=wal")
        self.cur4.execute("PRAGMA synchronous=OFF;")
        self.cur4.execute("CREATE TABLE IF NOT EXISTS path (document_num int, path text)")
        self.cur4.execute("CREATE INDEX IF NOT EXISTS path_IDX ON path (document_num)")

    def make_index(self, path=None, binary=None, contents=None):
        if path is not None or binary is not None:
            name, ext = os.path.splitext(path)
            if ext[1:] == 'ppt' or ext[1:] == 'xls' or ext[1:] == 'doc' or ext[1:] == 'pdf' or ext[1:] == 'hwp' or \
                    ext[1:] == 'pptx' or ext[1:] == 'xlsx' or ext[1:] == 'docx' or ext[1:] == 'txt':
                print(path, self.file_num)
                if binary is None:
                    m = MakeIndex(ext, path, self.file_num, self.conn1, self.cur1, None)
                else:
                    self.bin = io.BytesIO(binary)
                    m = MakeIndex(ext, path, self.file_num, self.conn1, self.cur1, self.bin)

                dic = m.data_info()
                for key, value in dic.items():
                    key = key.encode('utf-16', 'surrogatepass').decode('utf-16', errors='replace')
                    self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (key, value))
                self.cur2.execute("INSERT INTO path VALUES(?, ?)", (self.file_num, path))

                self.conn1.commit()
                self.conn2.commit()

                self.file_num += 1

        elif contents is not None:
            dic = {}
            start = 1
            contents = str(contents).upper()
            for i in range(len(contents)):
                compressed_start = VariableByteCode().VB_encode([start])
                compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                index_set = contents[i:i + 2]
                if index_set in dic:
                    dic[index_set] += str(start) + str(',')
                else:
                    dic[index_set] = str(start) + str(',')
                start += 1

            for key, value in dic.items():
                key = key.encode('utf-16', 'surrogatepass').decode('utf-16', errors='replace')
                self.cur3.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (key, value))
            self.cur4.execute()

            self.conn3.commit()

    def search(self, word):
        f = FindPath(self.index_db, self.path_db)
        for result in f.find_path(word):
            yield result
        for i in f.find_action(word):
            yield i


class MakeIndex:
    def __init__(self, ext, path, file_num, conn1, cur1, bin):
        self.ext = ext
        self.path = path
        self.file_num = file_num
        self.conn1 = conn1
        self.cur1 = cur1
        self.bin = bin
        self.file_name = str(path.split(os.sep)[-1])

    def data_info(self):
        dic = {}
        try:
            if self.ext[1:] == 'ppt' or self.ext[1:] == 'xls' or self.ext[1:] == 'doc' or self.ext[
                                                                                          1:] == 'pdf' or self.ext[
                                                                                                          1:] == 'hwp':
                headers = {"X-Tika-OCRLanguage": "eng", "X-Tika-OCRTimeout": "300"}
                if self.bin is None:
                    parsed = parser.from_file(self.path, xmlContent=False, requestOptions={'headers': headers})
                else:
                    parsed = parser.from_buffer(self.bin, xmlContent=False, requestOptions={'headers': headers})
                act = parsed['content']
                start = 1
                if act is not None:
                    buf = act.strip()
                else:
                    buf = ''
                buf = self.normalize_text(str(self.file_name) + buf)
                for i in range(len(buf)):
                    index_set = buf[i:i + 2]
                    if index_set == '  ' or index_set == ' ':
                        continue
                    else:
                        compressed_start = VariableByteCode().VB_encode([start])
                        compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                        if index_set in dic:
                            dic[index_set] += compressed_start
                        else:
                            dic[index_set] = compressed_file_num + compressed_start
                    start += 1
                    if len(dic[index_set]) > 100000:
                        index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16', errors='replace')
                        self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                          (index_set, dic[index_set]))
                        self.conn1.commit()
                        dic[index_set] = compressed_file_num
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'

            elif self.ext[1:] == 'pptx':
                if self.bin is None:
                    document = zipfile.ZipFile(self.path)
                else:
                    document = zipfile.ZipFile(self.bin)
                nums = []
                start = 1
                for d in document.namelist():
                    if d.startswith("ppt/slides/slide"):
                        nums.append(int(d[len("ppt/slides/slide"):-4]))
                s_format = "ppt/slides/slide%s.xml"
                slide_name_list = [s_format % x for x in sorted(nums)]
                buf = ''
                for slide in slide_name_list:
                    xml_content = document.read(slide)
                    tree = XML(xml_content)
                    if tree.tag == '{http://purl.oclc.org/ooxml/presentationml/main}sld':
                        NAMESPACE = '{http://purl.oclc.org/ooxml/drawingml/main}'
                    else:
                        NAMESPACE = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                    TEXT = NAMESPACE + 't'
                    for node in tree.iter(TEXT):
                        buf += str(node.text)
                buf = self.normalize_text(str(self.file_name) + buf)
                for i in range(len(buf)):
                    index_set = buf[i:i + 2]
                    if index_set == '  ' or index_set == ' ':
                        continue
                    else:
                        compressed_start = VariableByteCode().VB_encode([start])
                        compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                        if index_set in dic:
                            dic[index_set] += compressed_start
                        else:
                            dic[index_set] = compressed_file_num + compressed_start
                    start += 1
                    if len(dic[index_set]) > 100000:
                        index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16', errors='replace')
                        self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                          (index_set, dic[index_set]))
                        self.conn1.commit()
                        dic[index_set] = compressed_file_num
                document.close()
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'

            elif self.ext[1:] == 'xlsx':
                if self.bin is None:
                    workbook = xlrd.open_workbook(self.path, on_demand=True)
                else:
                    workbook = xlrd.open_workbook(filename=None, file_contents=self.bin.read(), on_demand=True)
                start = 1
                for sheet in workbook.sheet_names():
                    worksheet = workbook.sheet_by_name(sheet)
                    nrows = worksheet.nrows
                    ncols = worksheet.ncols
                    for row_num in range(nrows):
                        for col_num in range(ncols):
                            value = worksheet.cell_value(row_num, col_num)
                            if start == 1:
                                value = str(self.file_name) + value
                            buf = self.normalize_text(value)
                            for i in range(len(buf)):
                                index_set = buf[i:i + 2]
                                if index_set == '  ' or index_set == ' ':
                                    continue
                                else:
                                    compressed_start = VariableByteCode().VB_encode([start])
                                    compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                                    if index_set in dic:
                                        dic[index_set] += compressed_start
                                    else:
                                        dic[index_set] = compressed_file_num + compressed_start
                                start += 1
                                if len(dic[index_set]) > 100000:
                                    index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16',
                                                                                                   errors='replace')
                                    self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                                      (index_set, dic[index_set]))
                                    self.conn1.commit()
                                    dic[index_set] = compressed_file_num
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'

            elif self.ext[1:] == 'docx':
                if self.bin is None:
                    document = docx.Document(self.path)
                else:
                    document = docx.Document(self.bin)
                start = 1
                for x, paragraph in enumerate(document.paragraphs):
                    buf = self.normalize_text(paragraph.text)
                    if start == 1:
                        buf = str(self.file_name).upper() + buf
                    for i in range(len(buf)):
                        index_set = buf[i:i + 2]
                        if index_set == '  ' or index_set == ' ':
                            continue
                        else:
                            compressed_start = VariableByteCode().VB_encode([start])
                            compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                            if index_set in dic:
                                dic[index_set] += compressed_start
                            else:
                                dic[index_set] = compressed_file_num + compressed_start
                        start += 1
                        if len(dic[index_set]) > 100000:
                            index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16', errors='replace')
                            self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                              (index_set, dic[index_set]))
                            self.conn1.commit()
                            dic[index_set] = compressed_file_num
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'

            elif self.ext[1:] == 'txt':
                if self.bin is None:
                    check_f = open(self.path, 'rb')
                    check = check_f.readline()
                    detect = chardet.detect(check)
                    encoding = detect['encoding']
                    if encoding == 'ascii' or encoding == 'Windows-1254':
                        encoding = 'utf-8'
                    check_f.close()
                    f = open(self.path, 'r', encoding=encoding)
                    lines = f.readlines()
                    start = 1
                    for line in lines:
                        buf = self.normalize_text(line)
                        if start == 1:
                            buf = str(self.file_name).upper() + buf
                        for i in range(len(buf)):
                            index_set = buf[i:i + 2]
                            if index_set == '  ' or index_set == ' ':
                                continue
                            else:
                                compressed_start = VariableByteCode().VB_encode([start])
                                compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                                if index_set in dic:
                                    dic[index_set] += compressed_start
                                else:
                                    dic[index_set] = compressed_file_num + compressed_start
                            start += 1
                            if len(dic[index_set]) > 100000:
                                index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16',
                                                                                               errors='replace')
                                self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                                  (index_set, dic[index_set]))
                                self.conn1.commit()
                                dic[index_set] = compressed_file_num
                    for k, v in dic.items():
                        if dic[k][-1:] != b'\x00':
                            dic[k] = dic[k] + b'\x00'
                else:
                    check = self.bin.readline()
                    detect = chardet.detect(check)
                    encoding = detect['encoding']
                    if encoding == 'ascii' or encoding == 'Windows-1254':
                        encoding = 'utf-8'
                    lines = self.bin.readlines()
                    start = 1
                    for line in lines:
                        try:
                            buf = line.decode(encoding)
                            buf = self.normalize_text(buf)
                            if start == 1:
                                buf = str(self.file_name).upper() + buf
                            for i in range(len(buf)):
                                index_set = buf[i:i + 2]
                                if index_set == '  ' or index_set == ' ':
                                    continue
                                else:
                                    compressed_start = VariableByteCode().VB_encode([start])
                                    compressed_file_num = VariableByteCode().VB_encode([self.file_num])
                                    if index_set in dic:
                                        dic[index_set] += compressed_start
                                    else:
                                        dic[index_set] = compressed_file_num + compressed_start
                                start += 1
                                if len(dic[index_set]) > 100000:
                                    index_set = index_set.encode('utf-16', 'surrogatepass').decode('utf-16',
                                                                                                   errors='replace')
                                    self.cur1.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                                      (index_set, dic[index_set]))
                                    self.conn1.commit()
                                    dic[index_set] = compressed_file_num
                        except UnicodeDecodeError:
                            continue
                    for k, v in dic.items():
                        if dic[k][-1:] != b'\x00':
                            dic[k] = dic[k] + b'\x00'

        except:
            return dic

        return dic

    def normalize_text(self, buf):
        buf = str(buf).upper()
        buf = buf.replace('\r', ' ')
        buf = buf.replace('\n', ' ')
        buf = buf.replace('\t', ' ')
        return buf


class FindPath:
    def __init__(self, index_db, path_db):
        self.index_db = index_db
        self.path_db = path_db

    def word_ngrams(self, buf, n):
        output = []
        for i in range(len(buf) - n + 1):
            output.append(buf[i:i + n])
        return output

    def compress_concate(self, word):
        table_list = []
        word_list = self.word_ngrams(word, 2)
        conn = sqlite3.connect(self.index_db)
        cur = conn.cursor()
        for w in word_list:
            cur.execute("SELECT location FROM word_set WHERE word = ?", (w,))
            result = []
            for location in cur.fetchall():
                if type(location[0]) != str:
                    location = VariableByteCode().VB_decode(location[0])
                    if not result:
                        result += location
                    else:
                        if result[-1] != 0:
                            result = result + location[1:]
                        else:
                            result = result + location
            table_list.append(result)
        return table_list

    def make_table_entry(self, word):
        table_list = []
        for location in self.compress_concate(word):
            result = {}
            zero_index = [i for i, ele in enumerate(location) if ele == 0]
            p = 0
            for zero in zero_index:
                result[location[p]] = location[p + 1:zero]
                p = zero + 1
                if zero == zero_index[-1]:
                    break
            table_list.append(result)
        return table_list

    def make_input(self, dic):
        for k, v_list in dic.items():
            for v in v_list:
                yield k, v

    def binary_search(self, arr, target, low=None, high=None):
        try:
            low, high = low or 0, high or len(arr) - 1
            if low > high:
                return None
            mid = (low + high) // 2
            if arr[mid] > target:
                return self.binary_search(arr, target, low, mid)
            if arr[mid] == target:
                return target
            if arr[mid] < target:
                return self.binary_search(arr, target, mid + 1, high)
        except RecursionError:
            return None
        except IndexError:
            return None

    def find_wordset(self, word):
        table_list = self.make_table_entry(word)
        if len(table_list) >= 2:
            for i_k, i_v in self.make_input(table_list[0]):
                result = []
                for dic in table_list[1:]:
                    if i_k in dic:
                        v = i_v + 1
                        search_result = self.binary_search(sorted(dic[i_k]), v, 0, len(dic[i_k]))
                        if search_result is not None and search_result != -1:
                            result.append((i_k, i_v))
                            i_v = v
                        else:
                            break
                    else:
                        break
                yield result

        elif len(table_list) == 1:
            for i_k, i_v in self.make_input(table_list[0]):
                yield [(i_k, i_v)]

    def find_path(self, word):
        word = word.upper()
        conn = sqlite3.connect(self.path_db)
        cur = conn.cursor()

        if len(word) >= 2:
            word_list = self.word_ngrams(word, 2)
            for wordset in self.find_wordset(word):
                if len(wordset) == len(word_list) - 1 or len(word_list) == 1:
                    document_num = wordset[0][0]
                    cur.execute("SELECT path FROM path WHERE document_num = '" + str(document_num) + "'")
                    for path in cur.fetchall():
                        len_file_name = len(path[0].split(os.sep)[-1])
                        if wordset[0][1] < len_file_name:
                            yield path[0].replace('\\', '/'), None
                        else:
                            yield path[0].replace('\\', '/'), wordset[0][1] - len_file_name

        elif len(word) == 1:
            conn2 = sqlite3.connect(self.index_db)
            cur2 = conn2.cursor()
            cur2.execute("SELECT word FROM word_set WHERE word LIKE ?", ('%' + word + '%',))
            word_list = []
            for two_word in cur2.fetchall():
                word_list.append(two_word[0])
            word_list = OrderedDict.fromkeys(word_list)
            for w in word_list:
                for wordset in self.find_wordset(w):
                    document_num = wordset[0][0]
                    cur.execute("SELECT path FROM path WHERE document_num = '" + str(document_num) + "'")
                    for path in cur.fetchall():
                        len_file_name = len(path[0].split(os.sep)[-1])
                        if wordset[0][1] < len_file_name:
                            yield path[0].replace('\\', '/'), None
                        else:
                            yield path[0].replace('\\', '/'), wordset[0][1] - len_file_name

    def action_make_entry(self, word):
        table_list = []
        word_list = self.word_ngrams(word, 2)
        conn = sqlite3.connect(self.index_db)
        cur = conn.cursor()
        for w in word_list:
            dic = {}
            new_string = ''
            cur.execute("SELECT location FROM word_set WHERE word = ?", (w,))
            for location in cur.fetchall():
                if type(location[0]) == str:
                    new_string += location[0]
            dic[w] = self.string_to_list(new_string)
            table_list.append(dic)
        return table_list

    def string_to_list(self, string):
        result = []
        ten = ''
        for s in string:
            if s != ',':
                ten += s
            else:
                result.append(int(ten))
                ten = ''
        return result

    def action_find_wordset(self, word):
        w = self.word_ngrams(word, 2)
        table_list = self.action_make_entry(word)
        if len(table_list) >= 2:
            for i_k, i_v in self.make_input(table_list[0]):
                i = 1
                result = []
                for dic in table_list[i:]:
                    v = i_v + 1
                    search_result = self.binary_search(sorted(dic[w[i]]), v, 0, len(dic[w[i]]))
                    if search_result is not None and search_result != -1:
                        result.append((i_k, i_v))
                        i_v = v
                        i += 1
                    else:
                        break
                yield result

        elif len(table_list) == 1:
            for i_k, i_v in self.make_input(table_list[0]):
                yield [(i_k, i_v)]

    def find_action(self, string):
        string_up = string.upper()
        string_list = self.word_ngrams(string_up, 2)
        if len(string) >= 2:
            for wordset in self.action_find_wordset(string_up):
                if len(string_list) - 1 == len(wordset) or len(string_list) == 1:
                    yield string, wordset[0][1]

        elif len(string) == 1:
            conn2 = sqlite3.connect(self.index_db)
            cur2 = conn2.cursor()
            cur2.execute("SELECT word FROM word_set WHERE word LIKE ?", ('%' + string_up + '%',))
            word_list = []
            for two_word in cur2.fetchall():
                word_list.append(two_word[0])
            word_list = OrderedDict.fromkeys(word_list)
            for w in word_list:
                for wordset in self.action_find_wordset(w):
                    yield string, wordset[0][1]


class VariableByteCode:
    def __init__(self):
        pass

    @staticmethod
    def VB_encode_number(num):
        bytes_list = []
        while True:
            bytes_list.insert(0, num % 128)
            if num < 128:
                break
            num = num // 128
        bytes_list[-1] += 128
        return pack('%dB' % len(bytes_list), *bytes_list)

    def VB_encode(self, nums):
        bytes_list = []
        for number in nums:
            bytes_list.append(self.VB_encode_number(number))
        return b"".join(bytes_list)

    @staticmethod
    def VB_decode(bytestream):
        n = 0
        numbers = []
        bytestream = list(unpack('%dB' % len(bytestream), bytestream))
        zero_index = [i for i, ele in enumerate(bytestream) if ele == 0]
        for zero in zero_index:
            if bytestream[zero - 1] >= 128:
                bytestream[zero] = 128
        for byte in bytestream:
            if byte < 128:
                n = 128 * n + byte
            else:
                n = 128 * n + (byte - 128)
                numbers.append(n)
                n = 0
        return numbers


##########################################################################################
def get_file(drive):
    for root, dir, files in os.walk(drive):
        for file in files:
            full_path = os.path.join(root, file)
            try:
                f = open(full_path, 'rb')
                buf = f.read()
                yield full_path, buf
            except PermissionError:
                None
            except OSError:
                None


if __name__ == "__main__":
    from timeit import default_timer as timer
    from datetime import timedelta

    start = timer()

    folder = 'D:\\test'
    k = KeywordSearch(folder)
    '''for path, buf in get_file('C:\\'):
        k.make_index(path, buf)'''

    k.make_index(contents='www.google.com')
    '''for i in k.search('google'):
        print(i)'''

    end = timer()
    print(timedelta(seconds=end - start))