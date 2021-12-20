import os
import traceback
import chardet
import zipfile
import xlrd
import CompressInt
from tika import parser
import docx
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import tika
tika.initVM()


def normalize_text1(buf):
    buf = str(buf).upper()
    buf = buf.replace('\r', ' ')
    buf = buf.replace('\n', ' ')
    buf = buf.replace('\t', ' ')
    return buf

def make_dic(buf, insert_index):
    dic = {buf: insert_index}
    return dic

def ext(ext, full_path, file_num, conn, cur, bin=None):
    dic = {}
    file_name = full_path.split(os.sep)[-1]
    try:
        if ext[1:] == 'ppt' or ext[1:] == 'xls' or ext[1:] == 'doc' or ext[1:] == 'pdf' or ext[1:] == 'hwp':
            headers = {
                "X-Tika-OCRLanguage": "eng",
                "X-Tika-OCRTimeout": "300"
            }
            if bin is None:
                parsed = parser.from_file(full_path, xmlContent=False, requestOptions={'headers':headers, 'timeout':300})
            else:
                parsed = parser.from_buffer(bin, xmlContent=False, requestOptions={'headers': headers, 'timeout': 300})
            act = parsed['content']
            start = 1
            if act is not None:
                buf = act.strip()
            else:
                buf = ''
            buf = normalize_text1(file_name + buf)
            for i in range(len(buf)):
                index_set = buf[i:i + 2]
                if index_set == '  ' or index_set == ' ':
                    continue
                else:
                    compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                    compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                    result = make_dic(index_set, compressed_start)
                    if index_set in dic:
                        if dic[index_set][-1:] == b'\x00':
                            dic[index_set] += compressed_file_num
                        dic[index_set] += compressed_start
                    else:
                        result[index_set] = compressed_file_num + compressed_start
                        dic.update(result)
                start += 1
                if len(dic[index_set]) > 100000:
                    cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (index_set, dic[index_set]))
                    conn.commit()
                    dic[index_set] = compressed_file_num
            for k, v in dic.items():
                if dic[k][-1:] != b'\x00':
                    dic[k] = dic[k] + b'\x00'

        elif ext[1:] == 'pptx':
            if bin is None:
                document = zipfile.ZipFile(full_path)
            else:
                document = zipfile.ZipFile(bin)
            nums = []
            start = 1
            for d in document.namelist():
                if d.startswith("ppt/slides/slide"):
                    nums.append(int(d[len("ppt/slides/slide"):-4]))
            s_format = "ppt/slides/slide%s.xml"
            slide_name_list = [s_format % x for x in sorted(nums)]
            buf = ''
            for slide in slide_name_list:
                xml_content = document.read(slide)
                tree = XML(xml_content)
                if tree.tag == '{http://purl.oclc.org/ooxml/presentationml/main}sld':
                    NAMESPACE = '{http://purl.oclc.org/ooxml/drawingml/main}'
                else:
                    NAMESPACE = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                TEXT = NAMESPACE + 't'
                for node in tree.iter(TEXT):
                    buf += normalize_text1(node.text)
            buf = normalize_text1(file_name + buf)
            for i in range(len(buf)):
                index_set = buf[i:i + 2]
                if index_set == '  ' or index_set == ' ':
                    continue
                else:
                    compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                    compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                    result = make_dic(index_set, compressed_start)
                    if index_set in dic:
                        if dic[index_set][-1:] == b'\x00':
                            dic[index_set] += compressed_file_num
                        dic[index_set] += compressed_start
                    else:
                        result[index_set] = compressed_file_num + compressed_start
                        dic.update(result)
                start += 1
                if len(dic[index_set]) > 100000:
                    cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (index_set, dic[index_set]))
                    conn.commit()
                    dic[index_set] = compressed_file_num
            document.close()
            for k, v in dic.items():
                if dic[k][-1:] != b'\x00':
                    dic[k] = dic[k] + b'\x00'

        elif ext[1:] == 'xlsx':
            if bin is None:
                workbook = xlrd.open_workbook(full_path, on_demand=True)
            else:
                workbook = xlrd.open_workbook(filename=None, file_contents=bin.read(), on_demand=True)
            start = 1
            for sheet in workbook.sheet_names():
                worksheet = workbook.sheet_by_name(sheet)
                nrows = worksheet.nrows
                ncols = worksheet.ncols
                for row_num in range(nrows):
                    for col_num in range(ncols):
                        value = worksheet.cell_value(row_num, col_num)
                        buf = normalize_text1(value)
                        for i in range(len(buf)):
                            index_set = buf[i:i + 2]
                            if index_set == '  ' or index_set == ' ':
                                continue
                            else:
                                compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                                compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                                result = make_dic(index_set, compressed_start)
                                if index_set in dic:
                                    if dic[index_set][-1:] == b'\x00':
                                        dic[index_set] += compressed_file_num
                                    dic[index_set] += compressed_start
                                else:
                                    result[
                                        index_set] = compressed_file_num + compressed_start
                                    dic.update(result)
                            start += 1
                            if len(dic[index_set]) > 100000:
                                cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)",
                                            (index_set, dic[index_set]))
                                conn.commit()
                                dic[index_set] = compressed_file_num
            for k, v in dic.items():
                if dic[k][-1:] != b'\x00':
                    dic[k] = dic[k] + b'\x00'

        elif ext[1:] == 'docx':
            if bin is None:
                document = docx.Document(full_path)
            else:
                document = docx.Document(bin)
            start = 1
            for x, paragraph in enumerate(document.paragraphs):
                buf = normalize_text1(paragraph.text)
                for i in range(len(buf)):
                    index_set = buf[i:i + 2]
                    if index_set == '  ' or index_set == ' ':
                        continue
                    else:
                        compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                        compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                        result = make_dic(index_set, compressed_start)
                        if index_set in dic:
                            if dic[index_set][-1:] == b'\x00':
                                dic[index_set] += compressed_file_num
                            dic[index_set] += compressed_start
                        else:
                            result[
                                index_set] = compressed_file_num + compressed_start
                            dic.update(result)
                    start += 1
                    if len(dic[index_set]) > 100000:
                        cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (index_set, dic[index_set]))
                        conn.commit()
                        dic[index_set] = compressed_file_num
            for k, v in dic.items():
                if dic[k][-1:] != b'\x00':
                    dic[k] = dic[k] + b'\x00'

        elif ext[1:] == 'txt':
            if bin is None:
                check_f = open(full_path, 'rb')
                check = check_f.readline()
                detect = chardet.detect(check)
                encoding = detect['encoding']
                if encoding == 'ascii' or encoding == 'Windows-1254':
                    encoding = 'utf-8'
                check_f.close()
                f = open(full_path, 'r', encoding=encoding)
                lines = f.readlines()
                start = 1
                for line in lines:
                    buf = normalize_text1(line)
                    for i in range(len(buf)):
                        index_set = buf[i:i + 2]
                        if index_set == '  ' or index_set == ' ':
                            continue
                        else:
                            compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                            compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                            result = make_dic(index_set, compressed_start)
                            if index_set in dic:
                                if dic[index_set][-1:] == b'\x00':
                                    dic[index_set] += compressed_file_num
                                dic[index_set] += compressed_start
                            else:
                                result[index_set] = compressed_file_num + compressed_start
                                dic.update(result)
                        start += 1
                        if len(dic[index_set]) > 100000:
                            cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (index_set, dic[index_set]))
                            conn.commit()
                            dic[index_set] = compressed_file_num
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'
                f.close()

            else:
                check_f = bin
                check = check_f.readline()
                detect = chardet.detect(check)
                encoding = detect['encoding']
                if encoding == 'ascii' or encoding == 'Windows-1254':
                    encoding = 'utf-8'
                lines = bin.readlines()
                start = 1
                for line in lines:
                    buf = line.decode(encoding, errors='replace')
                    buf = normalize_text1(buf)
                    for i in range(len(buf)):
                        index_set = buf[i:i + 2]
                        if index_set == '  ' or index_set == ' ':
                            continue
                        else:
                            compressed_start = CompressInt.VariableByteCode().VB_encode([start])
                            compressed_file_num = CompressInt.VariableByteCode().VB_encode([file_num])
                            result = make_dic(index_set, compressed_start)
                            if index_set in dic:
                                if dic[index_set][-1:] == b'\x00':
                                    dic[index_set] += compressed_file_num
                                dic[index_set] += compressed_start
                            else:
                                result[index_set] = compressed_file_num + compressed_start
                                dic.update(result)
                        start += 1
                        if len(dic[index_set]) > 100000:
                            cur.execute("INSERT INTO word_set (word, location) VALUES (?, ?)", (index_set, dic[index_set]))
                            conn.commit()
                            dic[index_set] = compressed_file_num
                for k, v in dic.items():
                    if dic[k][-1:] != b'\x00':
                        dic[k] = dic[k] + b'\x00'

    except:
        error = traceback.format_exc()
        f = open(r'D:\test\errorlog.txt', 'a')
        f.write(full_path + '\n' + error + '\n')
        return dic

    return dic